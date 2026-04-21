# Copyright (c) 2026 Marco De Roni. All rights reserved.
# Licensed under the MIT License — see LICENSE file for details.

import os
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


SCORE_COLORS = {
    "RED":     RGBColor(0xC0, 0x00, 0x00),
    "YELLOW":  RGBColor(0xFF, 0xC0, 0x00),
    "GREEN":   RGBColor(0x00, 0x70, 0x00),
    "UNKNOWN": RGBColor(0x80, 0x80, 0x80),
}

SCORE_EMOJI = {
    "RED": "🔴",
    "YELLOW": "🟡",
    "GREEN": "🟢",
    "UNKNOWN": "⚪",
}


def add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading


def add_colored_paragraph(doc: Document, label: str, text: str, score: str):
    p = doc.add_paragraph()
    badge = p.add_run(f"{SCORE_EMOJI.get(score, '⚪')} [{score}]  ")
    badge.bold = True
    badge.font.color.rgb = SCORE_COLORS.get(score, SCORE_COLORS["UNKNOWN"])
    badge.font.size = Pt(10)
    lbl = p.add_run(f"{label}: ")
    lbl.bold = True
    lbl.font.size = Pt(10)
    txt = p.add_run(text)
    txt.font.size = Pt(10)


def generate_report(
    contract_name: str,
    metadata: dict,
    analysis: dict,
    output_dir: str,
    pii_summary: dict = None
) -> str:
    """Genera un report Word professionale."""

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    title = doc.add_heading("CONTRACT REVIEW REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Document: {contract_name}\n"
        f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # --- PII Redaction Summary ---
    if pii_summary and pii_summary.get("total_entities", 0) > 0:
        add_heading(doc, "Privacy & PII Redaction Summary", level=1)
        total = pii_summary.get("total_entities", 0)
        breakdown = pii_summary.get("breakdown", {})

        doc.add_paragraph(
            f"Before analysis, {total} sensitive entities were automatically "
            f"redacted and restored in this report."
        )

        if breakdown:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Entity Type"
            table.rows[0].cells[1].text = "Count"
            for entity_type, count in sorted(breakdown.items()):
                row = table.add_row().cells
                row[0].text = entity_type
                row[1].text = str(count)

        doc.add_paragraph()

    # --- Overall Risk ---
    overall = analysis.get("overall_score", "UNKNOWN")
    add_heading(doc, "Overall Risk Assessment", level=1)
    add_colored_paragraph(
        doc, "Overall Score",
        f"Contract risk level: {overall}",
        overall
    )
    doc.add_paragraph()

    # --- Metadata ---
    add_heading(doc, "Contract Metadata", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"

    metadata_labels = {
        "parties": "Parties",
        "effective_date": "Effective Date",
        "governing_law": "Governing Law",
        "jurisdiction": "Jurisdiction",
        "notice_period": "Notice Period",
        "duration": "Duration",
        "auto_renewal": "Auto-Renewal",
    }

    for key, label in metadata_labels.items():
        row = table.add_row().cells
        row[0].text = label
        row[1].text = metadata.get(key, "Not detected")

    doc.add_paragraph()

    # --- Missing Clauses ---
    add_heading(doc, "Missing Clauses", level=1)
    missing = analysis.get("missing_clauses", [])
    if missing:
        for clause in missing:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"⚠️  {clause} — not found in document")
            run.font.color.rgb = SCORE_COLORS["RED"]
            run.font.size = Pt(10)
    else:
        doc.add_paragraph("✅  All required clauses detected.")

    doc.add_paragraph()

    # --- Clause Risk Analysis ---
    add_heading(doc, "Clause Risk Analysis", level=1)
    findings = analysis.get("findings", {})

    if not findings:
        doc.add_paragraph("No risk patterns detected.")
    else:
        for category, items in findings.items():
            add_heading(doc, category.replace("_", " ").title(), level=2)
            for item in items:
                conf = item.get("confidence", "")
                add_colored_paragraph(
                    doc, item["score"],
                    f"{item['comment']} [{conf} confidence]",
                    item["score"]
                )
                if item.get("excerpt"):
                    excerpt_para = doc.add_paragraph()
                    excerpt_run = excerpt_para.add_run(item["excerpt"])
                    excerpt_run.font.size = Pt(9)
                    excerpt_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
                    excerpt_run.italic = True
                doc.add_paragraph()

    os.makedirs(output_dir, exist_ok=True)
    safe_name = contract_name.replace(" ", "_").replace("/", "_")
    output_path = os.path.join(output_dir, f"report_{safe_name}.docx")
    doc.save(output_path)
    return output_path